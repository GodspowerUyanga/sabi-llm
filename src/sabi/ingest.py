"""
Document text extraction for ingestion (PDF, DOCX, and plain formats).

These extractors run only at *ingest* time (when you add a document), never at
inference — so they do not affect the model's runtime RAM or the offline
guarantee. pypdf and python-docx are pure-Python and fully offline.
"""
from __future__ import annotations

from pathlib import Path

SUPPORTED_UPLOAD = {".pdf", ".docx", ".md", ".txt", ".csv", ".xlsx", ".xls"}


def extract_text(path: str | Path) -> str:
    """Extract plain text from a supported document."""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _from_pdf(path)
    if suffix == ".docx":
        return _from_docx(path)
    if suffix in (".xlsx", ".xls"):
        return _from_xlsx(path)
    # .md / .txt / .csv and anything else readable as text
    return path.read_text(encoding="utf-8", errors="replace")


def _from_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("Excel support needs openpyxl — run: pip install -r requirements.txt") from exc
    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts).strip()


def _from_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("PDF support needs pypdf — run: pip install -r requirements.txt") from exc
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n\n".join(pages).strip()


def _from_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("DOCX support needs python-docx — run: pip install -r requirements.txt") from exc
    document = docx.Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    # include table contents (common in business docs)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()
